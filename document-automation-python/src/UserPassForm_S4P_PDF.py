###CREATE DOCUMENT###
#!pip install python-docx

###DOCXtoPDF###
#!pip install docx2pdf
#!pip install docx


###CREATE DOCUMENT###
import os
import pandas as pd
import docx
from docx import Document
from docx.shared import Pt

###EMAIL###
import win32com.client as win32

###DOCXtoPDF###
from docx2pdf import convert
import time

cwd = os.getcwd()
path = os.path.join(cwd, "backendForms_Excel.xlsx")
path2 =  os.path.join(cwd, "backendForms_Vtemp.docx")
OUTPUT = os.path.join(cwd, "OUTPUT")

df = pd.read_excel(path)
doc = Document(path2)

Name = df['Userid']
Pass = df['Password']
Dep = df['Department']
Firstname = df['Firstname']
Sent_to = df['Email']

new_data = []

print('Start Process...')
for i in list(range(len(df))):
      
      print('Processing for ' + Name[i] + ' ...')
      #Department
      doc.tables[0].cell(0,1).paragraphs[0].text = Dep[i]
      new_data.append(doc.tables[0].cell(0,1).paragraphs[0])
      #UserID_0
      doc.tables[0].cell(1,1).paragraphs[0].text = Name[i]
      new_data.append(doc.tables[0].cell(1,1).paragraphs[0])
      #UserID_1
      doc.tables[1].cell(0,1).paragraphs[0].text = Name[i]
      new_data.append(doc.tables[1].cell(0,1).paragraphs[0])
      #Password
      doc.tables[1].cell(1,1).paragraphs[0].text = Pass[i]
      new_data.append(doc.tables[1].cell(1,1).paragraphs[0])

      for data in new_data:
        data.runs[0].font.name = 'TH SarabunPSK'
        data.runs[0].font.size = Pt(16)

      file_name_docx =  'backendForms_Vtemp_'+Name[i] + '.docx'
      doc.save(os.path.join(OUTPUT, file_name_docx))
      
      time.sleep(3)

      #####DOCtoPDF#####
      file_name_pdf =  'backendForms_'+Name[i] + '.pdf'
      inputFile = cwd + "\\OUTPUT\\" + file_name_docx
      outputFile = cwd + "\\OUTPUT\\" + file_name_pdf
      time.sleep(3)
      convert(inputFile, outputFile)

      #####SEND_MAIL#######  
      olApp = win32.Dispatch('Outlook.Application')
      olNS = olApp.GetNameSpace('MAPI')

      mailItem = olApp.CreateItem(0)
      mailItem.Subject = 'ขอนำส่ง Username และ Password ในระบบ CUERP S4 HANA สำหรับ USER: ' + Name[i]
      mailItem.BodyFormat = 1
      mailItem.HtmlBody = (
                            '<div style="font-family: TH Sarabun New, Arial, Tahoma; font-size:16pt;">'
                            'เรียนคุณ ' + Firstname[i] + '<br><br>&nbsp;&nbsp;&nbsp;&nbsp;'  # indent หน้า "ขอนำส่ง"
                            'ขอนำส่ง Username และ Password ในระบบ CUERP S4 HANA ตามเอกสารแนบค่ะ'
                            '<br><br>'
                            'ขอบคุณค่ะ'
                            '<br>'
                            'Hathaichanok Thumchirdchupong (Ing)'
                            '<br>'
                            'CUERP-80416'
                            '</div>'
                )
      #mailItem.To = 'dummy@email.ac.th'
      mailItem.To = Sent_to[i]
      mailItem.BCC = 'Nattaphon.K@chula.ac.th'
      mailItem.Attachments.Add(outputFile)

      mailItem.Display()
      #######################

print('Process Complete.')